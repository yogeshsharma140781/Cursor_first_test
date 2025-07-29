import os
from docx import Document
from docx.shared import RGBColor
import openai
from dotenv import load_dotenv

load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
client = openai.OpenAI(api_key=OPENAI_API_KEY)

def openai_translate_batch(texts, target_lang="English"):
    """Translate a batch of texts efficiently"""
    if not texts:
        return []
    
    # Combine all texts with clear separators
    combined_text = "\n\n---SEPARATOR---\n\n".join(texts)
    
    prompt = (
        f"Translate the following texts to {target_lang}. "
        f"CRITICAL PRESERVATION RULES - DO NOT TRANSLATE THESE: "
        f"1. Keep ALL web addresses EXACTLY as they are (www.example.com, http://..., https://...) "
        f"2. Keep ALL postal codes EXACTLY as written (1087 EM, 9560 AA, etc.) - These are location codes, NOT words to translate "
        f"3. Keep ALL street names, addresses, and city names UNCHANGED "
        f"4. Keep ALL phone numbers, email addresses, and reference numbers UNCHANGED "
        f"5. Keep ALL URLs and website paths UNCHANGED (including /path/to/page) "
        f"6. IMPORTANT: Letters like 'EM', 'AA', 'BB' in postal codes are NOT Dutch words - they are postal district codes "
        f"7. Translate day-of-week abbreviations (e.g., Mo, Di, Wo, Do, Vr) to their English equivalents (Mon, Tue, Wed, Thu, Fri). "
        f"8. IMPORTANT: Use '---SEPARATOR---' to separate each translation. Return exactly the same number of translations as input texts. "
        f"Only provide the translations, no explanations or additional text:\n\n"
        f"{combined_text}\n\nTranslation:"
    )
    
    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=2048,
            temperature=0.3
        )
        result = response.choices[0].message.content.strip()
        
        # Split the result back into individual translations
        translations = result.split("---SEPARATOR---")
        # Clean up each translation
        translations = [t.strip() for t in translations if t.strip()]
        
        # Ensure we have the same number of translations as input texts
        if len(translations) != len(texts):
            print(f"Warning: Expected {len(texts)} translations, got {len(translations)}")
            # Pad with original texts if needed
            while len(translations) < len(texts):
                translations.append(texts[len(translations)])
            # Truncate if we got too many
            translations = translations[:len(texts)]
        
        return translations
        
    except Exception as e:
        print(f"OpenAI translation error: {e}")
        return texts  # Return original texts on error

from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.parts.image import ImagePart
from docx.shared import Inches

def remove_all_paragraphs_from_header_footer(hf):
    # Remove all paragraph elements from a header/footer
    p_elements = hf._element.xpath('.//w:p')
    for p in p_elements:
        p.getparent().remove(p)

def copy_images_and_headers_footers(src_doc, dst_doc):
    """Copy all images and headers/footers from src_doc to dst_doc."""
    # Copy headers and footers
    for section_idx, section in enumerate(src_doc.sections):
        # Copy header
        src_header = section.header
        dst_header = dst_doc.sections[section_idx].header
        remove_all_paragraphs_from_header_footer(dst_header)
        for para in src_header.paragraphs:
            new_para = dst_header.add_paragraph()
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.font.size = run.font.size
                new_run.font.name = run.font.name
                if run.font.color.rgb:
                    new_run.font.color.rgb = run.font.color.rgb
            # Copy images in header
            for shape in para._element.xpath('.//w:drawing'):  # Inline images
                new_para._element.append(shape)
        # Copy footer
        src_footer = section.footer
        dst_footer = dst_doc.sections[section_idx].footer
        remove_all_paragraphs_from_header_footer(dst_footer)
        for para in src_footer.paragraphs:
            new_para = dst_footer.add_paragraph()
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.font.size = run.font.size
                new_run.font.name = run.font.name
                if run.font.color.rgb:
                    new_run.font.color.rgb = run.font.color.rgb
            for shape in para._element.xpath('.//w:drawing'):
                new_para._element.append(shape)
    # Copy inline images in body
    for i, para in enumerate(src_doc.paragraphs):
        for shape in para._element.xpath('.//w:drawing'):
            dst_doc.paragraphs[i]._element.append(shape)
    # Copy images in tables
    for t_idx, table in enumerate(src_doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    for shape in para._element.xpath('.//w:drawing'):
                        dst_doc.tables[t_idx].rows[r_idx].cells[c_idx].paragraphs[p_idx]._element.append(shape)

def translate_docx(input_path, output_path, target_lang="English"):
    """Translate all text in a DOCX file to the target language while preserving formatting and images"""
    try:
        # Load the document
        doc = Document(input_path)
        
        print(f"Processing document with {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables...")
        
        # Collect all text that needs translation
        text_elements = []
        element_info = []  # Store info about each element for later restoration
        
        # Collect paragraph texts
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                text_elements.append(paragraph.text)
                element_info.append({
                    'type': 'paragraph',
                    'index': i,
                    'runs': [(run.text, run.bold, run.italic, run.underline, run.font.size, run.font.name, run.font.color.rgb) for run in paragraph.runs]
                })
        
        # Collect table cell texts
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        if paragraph.text.strip():
                            text_elements.append(paragraph.text)
                            element_info.append({
                                'type': 'table_cell',
                                'table_idx': table_idx,
                                'row_idx': row_idx,
                                'cell_idx': cell_idx,
                                'para_idx': para_idx,
                                'runs': [(run.text, run.bold, run.italic, run.underline, run.font.size, run.font.name, run.font.color.rgb) for run in paragraph.runs]
                            })
        
        print(f"Collected {len(text_elements)} text elements for translation...")
        
        # Translate all texts in batches
        batch_size = 10  # Process 10 texts at a time
        all_translations = []
        
        for i in range(0, len(text_elements), batch_size):
            batch = text_elements[i:i+batch_size]
            print(f"Translating batch {i//batch_size + 1}/{(len(text_elements) + batch_size - 1)//batch_size} ({len(batch)} texts)...")
            batch_translations = openai_translate_batch(batch, target_lang)
            all_translations.extend(batch_translations)
        
        print(f"Translation completed. Applying translations with formatting preservation...")
        
        # Apply translations back to the document with formatting preservation
        translation_idx = 0
        
        # Apply to paragraphs
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                if translation_idx < len(all_translations):
                    translated_text = all_translations[translation_idx]
                    info = element_info[translation_idx]
                    
                    # Apply translation while preserving formatting
                    apply_translation_with_formatting(paragraph, translated_text, info['runs'])
                    translation_idx += 1
        
        # Apply to table cells
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        if paragraph.text.strip():
                            if translation_idx < len(all_translations):
                                translated_text = all_translations[translation_idx]
                                info = element_info[translation_idx]
                                
                                # Apply translation while preserving formatting
                                apply_translation_with_formatting(paragraph, translated_text, info['runs'])
                                translation_idx += 1
        
        # Save the translated document
        doc.save(output_path)
        print(f"Translation completed. Output saved to: {output_path}")
        
        # Copy images and headers/footers from original to translated doc
        print("Copying images and headers/footers...")
        src_doc = Document(input_path)
        dst_doc = Document(output_path)
        copy_images_and_headers_footers(src_doc, dst_doc)
        dst_doc.save(output_path)
        print("Images and headers/footers copied.")
        
    except Exception as e:
        print(f"Error translating document: {e}")
        raise

def apply_translation_with_formatting(paragraph, translated_text, original_runs_info):
    """Apply translated text to paragraph while preserving original formatting"""
    # Clear the paragraph
    paragraph.clear()
    
    # If we have original run formatting info, try to preserve it
    if original_runs_info and len(original_runs_info) > 0:
        # Simple approach: apply the first run's formatting to the entire translated text
        first_run_info = original_runs_info[0]
        run = paragraph.add_run(translated_text)
        
        # Apply formatting from the first run
        if first_run_info[1] is not None:  # bold
            run.bold = first_run_info[1]
        if first_run_info[2] is not None:  # italic
            run.italic = first_run_info[2]
        if first_run_info[3] is not None:  # underline
            run.underline = first_run_info[3]
        if first_run_info[4] is not None:  # font size
            run.font.size = first_run_info[4]
        if first_run_info[5] is not None:  # font name
            run.font.name = first_run_info[5]
        if first_run_info[6] is not None:  # font color
            run.font.color.rgb = first_run_info[6]
    else:
        # No formatting info available, just add the translated text
        paragraph.add_run(translated_text)

if __name__ == "__main__":
    import sys
    
    if len(sys.argv) != 3:
        print("Usage: python3 translate_docx_openai.py <input_docx> <output_docx>")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    if not os.path.exists(input_file):
        print(f"Error: Input file '{input_file}' not found.")
        sys.exit(1)
    
    print(f"Translating {input_file} to {output_file} with formatting and image preservation...")
    translate_docx(input_file, output_file) 