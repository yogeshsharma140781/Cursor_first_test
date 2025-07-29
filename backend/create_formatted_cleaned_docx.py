#!/usr/bin/env python3
"""
Create a new docx file with cleaned text while preserving the original layout and formatting
"""

import os
import json
from typing import List, Dict, Any
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime

class FormattedDocxCreator:
    def __init__(self):
        """Initialize the FormattedDocxCreator"""
        self.cleaned_data = None
        self.original_doc = None
        
    def load_cleaned_data(self, json_path: str) -> bool:
        """Load the cleaned text data from JSON file"""
        print(f"📋 Loading cleaned data from {json_path}...")
        
        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                self.cleaned_data = json.load(f)
            
            print(f"✅ Loaded data for {len(self.cleaned_data['paragraphs'])} paragraphs")
            return True
            
        except Exception as e:
            print(f"❌ Error loading cleaned data: {e}")
            return False
    
    def load_original_docx(self, docx_path: str) -> bool:
        """Load the original docx file to extract formatting"""
        print(f"📄 Loading original docx from {docx_path}...")
        
        try:
            self.original_doc = Document(docx_path)
            print(f"✅ Loaded original docx with {len(self.original_doc.paragraphs)} paragraphs")
            return True
            
        except Exception as e:
            print(f"❌ Error loading original docx: {e}")
            return False
    
    def get_paragraph_formatting(self, paragraph) -> Dict[str, Any]:
        """Extract formatting information from a paragraph"""
        formatting = {
            'alignment': paragraph.alignment,
            'style': paragraph.style.name if paragraph.style else None,
            'runs': []
        }
        
        for run in paragraph.runs:
            run_formatting = {
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'color': run.font.color.rgb if run.font.color.rgb else None
            }
            formatting['runs'].append(run_formatting)
        
        return formatting
    
    def apply_formatting_to_paragraph(self, new_paragraph, formatting: Dict[str, Any], cleaned_text: str):
        """Apply the original formatting to a new paragraph with cleaned text"""
        # Apply paragraph-level formatting
        if formatting['alignment']:
            new_paragraph.alignment = formatting['alignment']
        
        if formatting['style']:
            try:
                new_paragraph.style = formatting['style']
            except:
                pass  # Style might not exist in new document
        
        # Apply run-level formatting
        if formatting['runs']:
            # Split cleaned text into runs based on original run structure
            current_pos = 0
            for run_formatting in formatting['runs']:
                original_text = run_formatting['text']
                if original_text and current_pos < len(cleaned_text):
                    # Calculate how much text this run should contain
                    run_length = len(original_text)
                    run_text = cleaned_text[current_pos:current_pos + run_length]
                    
                    if run_text:
                        run = new_paragraph.add_run(run_text)
                        
                        # Apply formatting
                        if run_formatting['bold'] is not None:
                            run.bold = run_formatting['bold']
                        if run_formatting['italic'] is not None:
                            run.italic = run_formatting['italic']
                        if run_formatting['underline'] is not None:
                            run.underline = run_formatting['underline']
                        if run_formatting['font_name']:
                            run.font.name = run_formatting['font_name']
                        if run_formatting['font_size']:
                            run.font.size = run_formatting['font_size']
                        if run_formatting['color']:
                            run.font.color.rgb = run_formatting['color']
                    
                    current_pos += run_length
                else:
                    # If no original text, add a space to maintain formatting
                    run = new_paragraph.add_run(" ")
                    if run_formatting['bold'] is not None:
                        run.bold = run_formatting['bold']
                    if run_formatting['italic'] is not None:
                        run.italic = run_formatting['italic']
        else:
            # If no runs, just add the cleaned text
            new_paragraph.add_run(cleaned_text)
    
    def create_formatted_docx(self, output_path: str) -> bool:
        """Create a new docx file with cleaned text and original formatting"""
        print(f"📝 Creating formatted docx: {output_path}")
        
        if not self.cleaned_data or not self.original_doc:
            print("❌ Missing cleaned data or original document")
            return False
        
        try:
            # Create new document
            new_doc = Document()
            
            # Copy document-level properties
            new_doc.core_properties.title = self.original_doc.core_properties.title
            new_doc.core_properties.author = self.original_doc.core_properties.author
            new_doc.core_properties.subject = self.original_doc.core_properties.subject
            
            # Copy styles from original document
            for style in self.original_doc.styles:
                try:
                    if style.name not in new_doc.styles:
                        new_style = new_doc.styles.add_style(style.name, style.type)
                        # Copy style properties
                        if hasattr(style, 'font'):
                            new_style.font.name = style.font.name
                            new_style.font.size = style.font.size
                except:
                    pass  # Some styles might not be copyable
            
            # Process paragraphs
            cleaned_paragraphs = {p['index']: p for p in self.cleaned_data['paragraphs']}
            
            for i, original_para in enumerate(self.original_doc.paragraphs):
                # Get original formatting
                formatting = self.get_paragraph_formatting(original_para)
                
                # Check if we have cleaned text for this paragraph
                if i in cleaned_paragraphs:
                    cleaned_text = cleaned_paragraphs[i]['cleaned_text']
                    if cleaned_text:
                        # Create new paragraph with cleaned text
                        new_para = new_doc.add_paragraph()
                        self.apply_formatting_to_paragraph(new_para, formatting, cleaned_text)
                    else:
                        # Keep original text if no cleaned version
                        new_para = new_doc.add_paragraph()
                        self.apply_formatting_to_paragraph(new_para, formatting, original_para.text)
                else:
                    # Keep original paragraph if no cleaned version exists
                    new_para = new_doc.add_paragraph()
                    self.apply_formatting_to_paragraph(new_para, formatting, original_para.text)
            
            # Copy sections and page setup
            for section in self.original_doc.sections:
                new_section = new_doc.sections[0]  # New doc has one section by default
                new_section.page_width = section.page_width
                new_section.page_height = section.page_height
                new_section.left_margin = section.left_margin
                new_section.right_margin = section.right_margin
                new_section.top_margin = section.top_margin
                new_section.bottom_margin = section.bottom_margin
                new_section.header_distance = section.header_distance
                new_section.footer_distance = section.footer_distance
            
            # Save the document
            new_doc.save(output_path)
            print(f"✅ Formatted docx saved to: {output_path}")
            return True
            
        except Exception as e:
            print(f"❌ Error creating formatted docx: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def create_simple_formatted_docx(self, output_path: str) -> bool:
        """Create a simpler version that preserves basic formatting"""
        print(f"📝 Creating simple formatted docx: {output_path}")
        
        if not self.cleaned_data or not self.original_doc:
            print("❌ Missing cleaned data or original document")
            return False
        
        try:
            # Create new document
            new_doc = Document()
            
            # Copy document-level properties
            new_doc.core_properties.title = self.original_doc.core_properties.title
            new_doc.core_properties.author = self.original_doc.core_properties.author
            new_doc.core_properties.subject = self.original_doc.core_properties.subject
            
            # Process paragraphs with basic formatting preservation
            cleaned_paragraphs = {p['index']: p for p in self.cleaned_data['paragraphs']}
            
            for i, original_para in enumerate(self.original_doc.paragraphs):
                # Get basic formatting
                alignment = original_para.alignment
                style_name = original_para.style.name if original_para.style else None
                
                # Check if we have cleaned text for this paragraph
                if i in cleaned_paragraphs:
                    cleaned_text = cleaned_paragraphs[i]['cleaned_text']
                    if cleaned_text:
                        # Create new paragraph with cleaned text
                        new_para = new_doc.add_paragraph(cleaned_text)
                    else:
                        # Keep original text if no cleaned version
                        new_para = new_doc.add_paragraph(original_para.text)
                else:
                    # Keep original paragraph if no cleaned version exists
                    new_para = new_doc.add_paragraph(original_para.text)
                
                # Apply basic formatting
                if alignment:
                    new_para.alignment = alignment
                
                if style_name:
                    try:
                        new_para.style = style_name
                    except:
                        pass  # Style might not exist
            
            # Copy basic page setup
            if self.original_doc.sections:
                original_section = self.original_doc.sections[0]
                new_section = new_doc.sections[0]
                new_section.page_width = original_section.page_width
                new_section.page_height = original_section.page_height
                new_section.left_margin = original_section.left_margin
                new_section.right_margin = original_section.right_margin
                new_section.top_margin = original_section.top_margin
                new_section.bottom_margin = original_section.bottom_margin
            
            # Save the document
            new_doc.save(output_path)
            print(f"✅ Simple formatted docx saved to: {output_path}")
            return True
            
        except Exception as e:
            print(f"❌ Error creating simple formatted docx: {e}")
            import traceback
            traceback.print_exc()
            return False

def main():
    """Main function to create formatted cleaned docx"""
    print("🔧 Formatted DocX Creator")
    print("=" * 50)
    
    # File paths
    original_docx_path = "converted.docx"
    cleaned_json_path = "cleaned_output/cleaning_summary_20250704_133747.json"
    output_dir = "formatted_output"
    
    # Check if files exist
    if not os.path.exists(original_docx_path):
        print(f"❌ Original DocX file not found: {original_docx_path}")
        return
    
    if not os.path.exists(cleaned_json_path):
        print(f"❌ Cleaned JSON file not found: {cleaned_json_path}")
        return
    
    # Create output directory
    os.makedirs(output_dir, exist_ok=True)
    
    try:
        # Initialize creator
        creator = FormattedDocxCreator()
        
        # Load data
        if not creator.load_cleaned_data(cleaned_json_path):
            return
        
        if not creator.load_original_docx(original_docx_path):
            return
        
        # Create timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Create simple formatted version (more reliable)
        simple_output_path = os.path.join(output_dir, f"cleaned_formatted_simple_{timestamp}.docx")
        if creator.create_simple_formatted_docx(simple_output_path):
            print(f"✅ Simple formatted version created: {simple_output_path}")
        
        # Try to create full formatted version
        full_output_path = os.path.join(output_dir, f"cleaned_formatted_full_{timestamp}.docx")
        if creator.create_formatted_docx(full_output_path):
            print(f"✅ Full formatted version created: {full_output_path}")
        else:
            print("⚠️  Full formatting failed, simple version created instead")
        
        print(f"\n📁 Output files saved to: {output_dir}")
        
    except Exception as e:
        print(f"❌ Error during processing: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main() 