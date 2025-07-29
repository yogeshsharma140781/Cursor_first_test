#!/usr/bin/env python3
"""
Clean up garbled text in converted.docx using OpenAI Vision API
by comparing with the original scanned.png image
"""

import os
import base64
from typing import List, Dict, Any
from docx import Document
from docx.shared import Inches
from openai import OpenAI
import json
import re
from datetime import datetime

class DocxCleaner:
    def __init__(self):
        """Initialize the DocxCleaner with OpenAI client"""
        self.api_key = os.getenv('OPENAI_API_KEY')
        if not self.api_key:
            raise ValueError("❌ No OpenAI API key found in environment variables")
        
        self.client = OpenAI(api_key=self.api_key)
        print("✅ OpenAI client initialized successfully")
    
    def encode_image(self, image_path: str) -> str:
        """Encode image to base64 string"""
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')
    
    def extract_text_from_docx(self, docx_path: str) -> List[Dict[str, Any]]:
        """Extract text from docx file paragraph by paragraph"""
        print(f"📖 Extracting text from {docx_path}...")
        
        doc = Document(docx_path)
        paragraphs = []
        
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():  # Only include non-empty paragraphs
                paragraphs.append({
                    'index': i,
                    'original_text': para.text.strip(),
                    'cleaned_text': None,
                    'confidence': None
                })
        
        print(f"✅ Extracted {len(paragraphs)} paragraphs from docx")
        return paragraphs
    
    def analyze_image_with_vision(self, image_path: str) -> str:
        """Use OpenAI Vision API to extract and understand the original text"""
        print("🔍 Analyzing original scanned image with OpenAI Vision...")
        
        base64_image = self.encode_image(image_path)
        
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": """Please analyze this scanned document image and extract all the text content. 
                                
                                Focus on:
                                1. Identifying all paragraphs and their content
                                2. Understanding the structure and layout
                                3. Recognizing any headers, subheadings, or special formatting
                                4. Noting any tables, lists, or structured content
                                
                                Return the extracted text in a structured format that preserves the document's organization."""
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{base64_image}"
                                }
                            }
                        ]
                    }
                ],
                max_tokens=2000,
                temperature=0.1
            )
            
            original_text = response.choices[0].message.content
            print("✅ Successfully analyzed original image")
            return original_text
            
        except Exception as e:
            print(f"❌ Error analyzing image: {e}")
            return None
    
    def clean_paragraph_text(self, garbled_text: str, original_context: str, paragraph_index: int) -> Dict[str, Any]:
        """Clean up a single paragraph using OpenAI"""
        print(f"🧹 Cleaning paragraph {paragraph_index + 1}...")
        
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": """You are an expert text restoration specialist. Your task is to clean up garbled/corrupted text from OCR or document conversion by comparing it with the original document context.

Rules:
1. Fix obvious OCR errors (e.g., 'rn' -> 'm', 'cl' -> 'd', etc.)
2. Restore proper spacing and punctuation
3. Correct word boundaries that were merged or split incorrectly
4. Maintain the original meaning and intent
5. Don't add information that wasn't in the original
6. If uncertain about a word, use context clues from the original document
7. Return only the cleaned text, no explanations

Focus on making the text readable and accurate while preserving the original structure and meaning."""
                    },
                    {
                        "role": "user",
                        "content": f"""Please clean up this garbled text based on the original document context:

GARBLED TEXT:
{garbled_text}

ORIGINAL DOCUMENT CONTEXT:
{original_context}

PARAGRAPH INDEX: {paragraph_index + 1}

Please return only the cleaned, corrected text."""
                    }
                ],
                temperature=0.1,
                max_tokens=800
            )
            
            cleaned_text = response.choices[0].message.content.strip()
            
            # Calculate a simple confidence score based on text similarity
            confidence = self.calculate_confidence(garbled_text, cleaned_text)
            
            return {
                'cleaned_text': cleaned_text,
                'confidence': confidence,
                'processing_successful': True
            }
            
        except Exception as e:
            print(f"❌ Error cleaning paragraph {paragraph_index + 1}: {e}")
            return {
                'cleaned_text': garbled_text,  # Return original if cleaning fails
                'confidence': 0.0,
                'processing_successful': False
            }
    
    def calculate_confidence(self, original: str, cleaned: str) -> float:
        """Calculate a simple confidence score for the cleaning"""
        # Simple metrics: length similarity, common words, etc.
        orig_words = set(original.lower().split())
        clean_words = set(cleaned.lower().split())
        
        if len(orig_words) == 0:
            return 0.0
        
        # Jaccard similarity
        intersection = len(orig_words.intersection(clean_words))
        union = len(orig_words.union(clean_words))
        
        if union == 0:
            return 0.0
        
        similarity = intersection / union
        
        # Also consider length similarity
        length_ratio = min(len(original), len(cleaned)) / max(len(original), len(cleaned))
        
        # Combined score
        confidence = (similarity * 0.7) + (length_ratio * 0.3)
        return round(confidence, 2)
    
    def create_cleaned_docx(self, paragraphs: List[Dict[str, Any]], output_path: str):
        """Create a new cleaned docx file"""
        print(f"📝 Creating cleaned docx file: {output_path}")
        
        doc = Document()
        
        # Add a title
        title = doc.add_heading('Cleaned Document', 0)
        
        # Add metadata
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("Cleaned using OpenAI Vision API")
        doc.add_paragraph("-" * 50)
        
        # Add cleaned paragraphs
        for para in paragraphs:
            if para['cleaned_text']:
                p = doc.add_paragraph(para['cleaned_text'])
                
                # Add confidence info as a comment-like paragraph
                if para['confidence'] is not None:
                    confidence_para = doc.add_paragraph(f"[Confidence: {para['confidence']:.2f}]")
                    confidence_para.style = 'Intense Quote'
        
        doc.save(output_path)
        print(f"✅ Cleaned docx saved to: {output_path}")
    
    def create_comparison_report(self, paragraphs: List[Dict[str, Any]], output_path: str):
        """Create a detailed comparison report"""
        print(f"📊 Creating comparison report: {output_path}")
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Text Cleaning Comparison Report', 0)
        
        # Summary
        doc.add_heading('Summary', level=1)
        total_paragraphs = len(paragraphs)
        successful_cleanings = sum(1 for p in paragraphs if p.get('processing_successful', False))
        avg_confidence = sum(p.get('confidence', 0) for p in paragraphs) / total_paragraphs if total_paragraphs > 0 else 0
        
        doc.add_paragraph(f"Total paragraphs processed: {total_paragraphs}")
        doc.add_paragraph(f"Successfully cleaned: {successful_cleanings}")
        doc.add_paragraph(f"Average confidence: {avg_confidence:.2f}")
        doc.add_paragraph(f"Processing date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Detailed comparisons
        doc.add_heading('Detailed Comparisons', level=1)
        
        for i, para in enumerate(paragraphs):
            doc.add_heading(f'Paragraph {i + 1}', level=2)
            
            # Original text
            doc.add_paragraph("Original (Garbled):")
            original_p = doc.add_paragraph(para['original_text'])
            original_p.style = 'Intense Quote'
            
            # Cleaned text
            doc.add_paragraph("Cleaned:")
            cleaned_p = doc.add_paragraph(para['cleaned_text'] or 'Failed to clean')
            cleaned_p.style = 'Quote'
            
            # Confidence and status
            confidence = para.get('confidence', 0)
            status = "✅ Success" if para.get('processing_successful', False) else "❌ Failed"
            doc.add_paragraph(f"Status: {status} | Confidence: {confidence:.2f}")
            
            doc.add_paragraph("-" * 40)
        
        doc.save(output_path)
        print(f"✅ Comparison report saved to: {output_path}")
    
    def process_docx(self, docx_path: str, image_path: str, output_dir: str = "cleaned_output"):
        """Main process to clean up the docx file"""
        print("🚀 Starting docx cleaning process...")
        
        # Create output directory
        os.makedirs(output_dir, exist_ok=True)
        
        # Step 1: Extract text from docx
        paragraphs = self.extract_text_from_docx(docx_path)
        
        if not paragraphs:
            print("❌ No text found in docx file")
            return
        
        # Step 2: Analyze original image
        original_context = self.analyze_image_with_vision(image_path)
        
        if not original_context:
            print("❌ Could not analyze original image")
            return
        
        print(f"📄 Original document context extracted:\n{original_context[:200]}...")
        
        # Step 3: Clean each paragraph
        print(f"🧹 Starting to clean {len(paragraphs)} paragraphs...")
        
        for para in paragraphs:
            result = self.clean_paragraph_text(
                para['original_text'], 
                original_context, 
                para['index']
            )
            
            para.update(result)
        
        # Step 4: Create outputs
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Create cleaned docx
        cleaned_docx_path = os.path.join(output_dir, f"cleaned_document_{timestamp}.docx")
        self.create_cleaned_docx(paragraphs, cleaned_docx_path)
        
        # Create comparison report
        report_path = os.path.join(output_dir, f"cleaning_report_{timestamp}.docx")
        self.create_comparison_report(paragraphs, report_path)
        
        # Create JSON summary
        json_path = os.path.join(output_dir, f"cleaning_summary_{timestamp}.json")
        summary = {
            'processing_date': datetime.now().isoformat(),
            'total_paragraphs': len(paragraphs),
            'successful_cleanings': sum(1 for p in paragraphs if p.get('processing_successful', False)),
            'average_confidence': sum(p.get('confidence', 0) for p in paragraphs) / len(paragraphs),
            'paragraphs': paragraphs
        }
        
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(summary, f, indent=2, ensure_ascii=False)
        
        print(f"✅ Processing complete! Files saved to: {output_dir}")
        print(f"📄 Cleaned document: {cleaned_docx_path}")
        print(f"📊 Comparison report: {report_path}")
        print(f"📋 JSON summary: {json_path}")
        
        return summary

def main():
    """Main function to run the cleaning process"""
    print("🔧 DocX Cleaner with OpenAI Vision")
    print("=" * 50)
    
    # File paths
    docx_path = "converted.docx"
    image_path = "scanned.png"
    output_dir = "cleaned_output"
    
    # Check if files exist
    if not os.path.exists(docx_path):
        print(f"❌ DocX file not found: {docx_path}")
        return
    
    if not os.path.exists(image_path):
        print(f"❌ Image file not found: {image_path}")
        return
    
    try:
        # Initialize cleaner
        cleaner = DocxCleaner()
        
        # Process the docx file
        summary = cleaner.process_docx(docx_path, image_path, output_dir)
        
        # Print final summary
        print("\n📈 Final Summary:")
        print(f"Total paragraphs: {summary['total_paragraphs']}")
        print(f"Successfully cleaned: {summary['successful_cleanings']}")
        print(f"Average confidence: {summary['average_confidence']:.2f}")
        print(f"Success rate: {(summary['successful_cleanings'] / summary['total_paragraphs'] * 100):.1f}%")
        
    except Exception as e:
        print(f"❌ Error during processing: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main() 