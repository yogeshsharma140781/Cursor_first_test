#!/usr/bin/env python3

import os
import sys
from pathlib import Path
from docx import Document
from googletrans import Translator
import logging

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('doc_translator.log'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

class DocTranslator:
    def __init__(self, input_docx_path, output_pdf_path, target_lang='en'):
        """
        Initialize the translator with input and output paths.
        
        Args:
            input_docx_path (str): Path to the input DOCX file
            output_pdf_path (str): Path where the translated PDF will be saved
            target_lang (str): Target language code (default: 'en' for English)
        """
        self.input_docx_path = Path(input_docx_path)
        self.output_pdf_path = Path(output_pdf_path)
        self.target_lang = target_lang
        self.translator = Translator()

    def translate_docx(self):
        """Translate the content of the DOCX file while preserving text block structure and alignment."""
        try:
            logger.info("Starting DOCX translation")
            doc = Document(self.input_docx_path)
            
            def translate_paragraph(paragraph):
                if not paragraph.text.strip():
                    return

                # Store original paragraph properties
                original_alignment = paragraph.alignment
                original_style = paragraph.style
                original_paragraph_format = paragraph.paragraph_format
                original_spacing_before = original_paragraph_format.space_before
                original_spacing_after = original_paragraph_format.space_after
                original_line_spacing = original_paragraph_format.line_spacing
                original_left_indent = original_paragraph_format.left_indent
                original_right_indent = original_paragraph_format.right_indent
                original_first_line_indent = original_paragraph_format.first_line_indent

                # Group runs by their formatting to maintain style consistency and preserve spaces
                current_text = ""
                current_format = None
                formatted_runs = []

                for idx, run in enumerate(paragraph.runs):
                    if not run.text.strip():
                        # Still preserve spaces if present
                        if run.text and current_format is not None:
                            current_text += run.text
                        continue

                    run_format = {
                        'font_name': run.font.name,
                        'font_size': run.font.size,
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'color': run.font.color.rgb if run.font.color else None
                    }

                    if current_format == run_format:
                        # Add a space if the previous text does not end with space and this run does not start with punctuation
                        if current_text and not current_text[-1].isspace() and not run.text[0] in ',.;:!?':
                            current_text += ' '
                        current_text += run.text
                    else:
                        if current_text:
                            formatted_runs.append((current_text, current_format))
                        current_text = run.text
                        current_format = run_format

                if current_text:
                    formatted_runs.append((current_text, current_format))

                # Clear paragraph and translate each formatted group
                paragraph.clear()
                for text, format_dict in formatted_runs:
                    try:
                        translated = self.translator.translate(text, dest=self.target_lang)
                        new_run = paragraph.add_run(translated.text)
                        
                        # Apply original formatting
                        if format_dict['font_name']:
                            new_run.font.name = format_dict['font_name']
                        if format_dict['font_size']:
                            new_run.font.size = format_dict['font_size']
                        new_run.bold = format_dict['bold']
                        new_run.italic = format_dict['italic']
                        new_run.underline = format_dict['underline']
                        if format_dict['color']:
                            new_run.font.color.rgb = format_dict['color']
                    except Exception as e:
                        logger.warning(f"Error translating text block: {str(e)}")
                        # If translation fails, keep original text
                        new_run = paragraph.add_run(text)
                        # Apply original formatting
                        if format_dict['font_name']:
                            new_run.font.name = format_dict['font_name']
                        if format_dict['font_size']:
                            new_run.font.size = format_dict['font_size']
                        new_run.bold = format_dict['bold']
                        new_run.italic = format_dict['italic']
                        new_run.underline = format_dict['underline']
                        if format_dict['color']:
                            new_run.font.color.rgb = format_dict['color']

                # Restore paragraph properties
                paragraph.alignment = original_alignment
                paragraph.style = original_style
                paragraph.paragraph_format.space_before = original_spacing_before
                paragraph.paragraph_format.space_after = original_spacing_after
                paragraph.paragraph_format.line_spacing = original_line_spacing
                paragraph.paragraph_format.left_indent = original_left_indent
                paragraph.paragraph_format.right_indent = original_right_indent
                paragraph.paragraph_format.first_line_indent = original_first_line_indent

            # Translate paragraphs
            for paragraph in doc.paragraphs:
                translate_paragraph(paragraph)

            # Translate tables while preserving structure
            for table in doc.tables:
                # Store table properties
                table_style = table.style
                table_alignment = table.alignment
                
                for row in table.rows:
                    # Store row properties
                    row_height = row.height
                    
                    for cell in row.cells:
                        # Store cell properties
                        cell_width = cell.width
                        cell_vertical_alignment = cell.vertical_alignment
                        
                        for paragraph in cell.paragraphs:
                            translate_paragraph(paragraph)
                            
                        # Restore cell properties
                        cell.width = cell_width
                        cell.vertical_alignment = cell_vertical_alignment
                    
                    # Restore row properties
                    row.height = row_height
                
                # Restore table properties
                table.style = table_style
                table.alignment = table_alignment

            # Translate headers and footers
            for section in doc.sections:
                # Header
                header = section.header
                for paragraph in header.paragraphs:
                    translate_paragraph(paragraph)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                translate_paragraph(paragraph)
                # Footer
                footer = section.footer
                for paragraph in footer.paragraphs:
                    translate_paragraph(paragraph)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                translate_paragraph(paragraph)

            # Save the translated document
            translated_docx_path = self.input_docx_path.with_name(
                f"{self.input_docx_path.stem}_translated.docx"
            )
            doc.save(translated_docx_path)
            logger.info("DOCX translation completed successfully")
            return translated_docx_path
        except Exception as e:
            logger.error(f"Error during DOCX translation: {str(e)}")
            return None

    def docx_to_pdf(self, translated_docx_path):
        """Convert the translated DOCX back to PDF using LibreOffice."""
        try:
            logger.info(f"Converting translated DOCX to PDF: {translated_docx_path}")
            # Use LibreOffice's soffice command to convert DOCX to PDF
            cmd = f'soffice --headless --convert-to pdf --outdir "{self.output_pdf_path.parent}" "{translated_docx_path}"'
            result = os.system(cmd)
            
            if result == 0:
                # LibreOffice saves the PDF with the same name but .pdf extension
                expected_pdf = translated_docx_path.with_suffix('.pdf')
                if expected_pdf.exists():
                    # Move the generated PDF to the desired output path
                    expected_pdf.rename(self.output_pdf_path)
                    logger.info("DOCX to PDF conversion completed successfully")
                    return True
                else:
                    logger.error("PDF file was not generated by LibreOffice")
                    return False
            else:
                logger.error(f"LibreOffice conversion failed with exit code {result}")
                return False
        except Exception as e:
            logger.error(f"Error converting DOCX to PDF: {str(e)}")
            return False

    def translate(self):
        """Execute the complete translation workflow."""
        try:
            # Step 1: Translate the DOCX
            translated_docx_path = self.translate_docx()
            if not translated_docx_path:
                return False

            # Step 2: Convert to PDF
            if not self.docx_to_pdf(translated_docx_path):
                return False
            
            logger.info("Translation workflow completed successfully")
            return True
        except Exception as e:
            logger.error(f"Error in translation workflow: {str(e)}")
            return False

def main():
    if len(sys.argv) < 3:
        print("Usage: python doc_based_translator.py <input_docx_path> <output_pdf_path> [target_lang]")
        sys.exit(1)

    input_docx = sys.argv[1]
    output_pdf = sys.argv[2]
    target_lang = sys.argv[3] if len(sys.argv) > 3 else 'en'

    translator = DocTranslator(input_docx, output_pdf, target_lang)
    success = translator.translate()

    if success:
        print(f"Translation completed successfully. Output saved to: {output_pdf}")
    else:
        print("Translation failed. Check the log file for details.")
        sys.exit(1)

if __name__ == "__main__":
    main() 